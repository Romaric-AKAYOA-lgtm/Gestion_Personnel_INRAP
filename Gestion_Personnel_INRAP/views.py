import matplotlib.pyplot as plt
import base64
from django.db.models.functions import Lower
from io import BytesIO
from django.shortcuts import render, redirect
from django.db.models import Q, Count, F
from django.utils.timezone import now
from datetime import datetime, timedelta, timezone
from django.utils import timezone
from datetime import timedelta
from Activation.models import Activation
from Conge.models import Conge
from Employee.models import Employee
from Employee.views import get_username_from_session
from OrganizationalUnit.models import OrganizationalUnit
from RespensableOrganisationUnite.models import ResponsableOrganisationUnite
from Stagiaire.models import Stagiaire

def home_view(request):
    # Vérification de l'activation
    activation = Activation.objects.first()
    if not activation or not activation.is_valid():
        return redirect("Activation:activation_page")

    username = get_username_from_session(request)

    # Assurez-vous que le nom d'utilisateur est disponible dans la session
    if not username:
        return redirect('login')  # Redirige vers la page de connexion si pas de nom d'utilisateur dans la session


    # Obtenir la date d'aujourd'hui
    today = timezone.now().date()

    # Récupérer l'année actuelle
    current_year = today.year

    # Calculer la date six mois en arrière
    six_months_ago = today - timedelta(days=180)
    # Recherche d'employés par nom ou fonction
    recherche = request.GET.get("search", "")
    employes = Employee.objects.all()
    if recherche:
        employes = employes.filter(Q(last_name__icontains=recherche) | Q(function__icontains=recherche))

    # Filtrer les employés récents (moins de 6 mois)
    employes_recents = employes.filter(start_date__gte=six_months_ago)

    # Employés en congé
    
    # **Employés en congé**
    # Employés de retour de congé (date_fin < aujourd'hui)
    employes_retour_conge = Conge.objects.filter(
        date_fin__lt=today  # Le congé s'est terminé avant aujourd'hui
    ).exclude(date_debut__isnull=True).exclude(date_fin__isnull=True).distinct()

    # Employés actuellement en congé (date_fin >= aujourd'hui)
    employes_en_conge = Conge.objects.filter(
        date_fin__gte=today  # Le congé continue ou termine après ou aujourd'hui
    ).exclude(date_debut__isnull=True).exclude(date_fin__isnull=True).distinct()

    # Calcul de l'âge des employés
    ages_employes = [
        {
            "employee": emp,
            "age": (current_year - emp.date_of_birth.year) - ((today.month, today.day) < (emp.date_of_birth.month, emp.date_of_birth.day)) if emp.date_of_birth else None
        }
        for emp in employes
    ]

    moins_de_30_ans = sum(1 for item in ages_employes if item["age"] is not None and item["age"] < 30)
    de_30_a_50_ans = sum(1 for item in ages_employes if item["age"] is not None and 30 <= item["age"] < 50)
    de_50_a_60_ans = sum(1 for item in ages_employes if item["age"] is not None and 50 <= item["age"] < 60)

    employes_partant_en_retraite = [emp["employee"] for emp in ages_employes if emp["age"] == 60]
    employes_retraites = [emp["employee"] for emp in ages_employes if emp["age"] is not None and emp["age"] > 60]

    # Récupération des responsables d'unités organisationnelles et comptage des employés par unité
    responsables_unites = ResponsableOrganisationUnite.objects.annotate(
        nombre_employes=Count('responsable')  # Comptage des employés (responsables) associés à chaque unité
    )

    # Récupérer la date actuelle du système
    current_date = timezone.now().date()

    # Filtrer les responsables dont la date de fin est None ou supérieure à la date système
    responsables_Organisation_unites = ResponsableOrganisationUnite.objects.filter(
        Q(date_fin__isnull=True) | Q(date_fin__gte=current_date)  # Date de fin = None ou supérieure à la date actuelle
    )

    # Utilisation de la jointure pour compter les hommes et les femmes
    employes_par_responsable = ResponsableOrganisationUnite.objects.annotate(
        male_count=Count('responsable', filter=Q(responsable__sexe="Masculin")),
        female_count=Count('responsable', filter=Q(responsable__sexe="Féminin"))
    )

    # Calcul du total des hommes et femmes dans tous les responsables
    employes_hommes = employes_par_responsable.aggregate(total_males=Count('responsable', filter=Q(responsable__sexe="Masculin")))["total_males"]
    employes_femmes = employes_par_responsable.aggregate(total_females=Count('responsable', filter=Q(responsable__sexe="Féminin")))["total_females"]

    # Récupération des données sur le sexe des responsables
    donnees_sexe_responsable = get_responsible_gender_data(employes_par_responsable)
    graphique_sexe = create_chart_base64(donnees_sexe_responsable, "Progression des hommes et femmes", "Sexe", "Nombre")

    # Progression des employés
    progression_employes = ResponsableOrganisationUnite.objects.annotate(
        employee_count=Count('responsable')  # Nombre d'employés dans chaque unité organisationnelle
    )

    # Progression de la retraite
    progression_retraite = get_retirement_progression(ages_employes)
    graphique_retraite = create_chart_base64(progression_retraite, "Progression des employés partant à la retraite", "Catégorie", "Nombre")

    # Répartition des hommes et des femmes par unité organisationnelle
    repartition_sexe_par_unite = get_gender_distribution_per_unit()
    graphique_repartition = create_chart_base64(repartition_sexe_par_unite, "Répartition des hommes et femmes par unité", "Unité Organisationnelle", "Nombre", True)

    return render(request, "home.html", {
        "username": username,
        "employes": employes,
        "moins_de_30_ans": moins_de_30_ans,
        "de_30_a_50_ans": de_30_a_50_ans,
        "de_50_a_60_ans": de_50_a_60_ans,
        "employes_recents": employes_recents,
        "employes_retour_conge": employes_retour_conge,
        "employes_en_conge": employes_en_conge,
        "employes_partant_en_retraite": employes_partant_en_retraite,
        "employes_retraites": employes_retraites,
        "responsables_unites": responsables_unites,
        "organisation_unite": OrganizationalUnit.objects.all().order_by('parent'),
        "employes_par_responsable": employes_par_responsable,
        "progression_employes": progression_employes,
        "employes_hommes": employes_hommes,
        "employes_femmes": employes_femmes,
        "graphique_sexe": graphique_sexe,
        "graphique_retraite": graphique_retraite,
        "graphique_repartition": graphique_repartition,
        'responsables_Organisation_unites':responsables_Organisation_unites,
    })

def create_chart_base64(data, title, xlabel, ylabel, is_bar=True):
    fig, ax = plt.subplots()

    # Conversion des données en listes
    keys = list(data.keys())  # Assurez-vous que 'keys' est une liste de chaînes
    values = list(data.values())  # Assurez-vous que 'values' est une liste

    # Vérification si les valeurs sont des dictionnaires et les convertir
    if isinstance(values[0], dict):
        # Si les valeurs sont des dictionnaires, transformez-les en listes de valeurs
        values = [sum(val.values()) for val in values]
    # Assurez-vous que les clés et les valeurs sont sous forme de types hachables
    if is_bar:
        ax.bar(keys, values)
    else:
        ax.pie(values, labels=keys, autopct="%1.1f%%")
    
    ax.set_title(title)
    ax.set_xlabel(xlabel)
    ax.set_ylabel(ylabel)

    buf = BytesIO()
    fig.savefig(buf, format="png")
    buf.seek(0)
    image_base64 = base64.b64encode(buf.getvalue()).decode("utf-8")
    return f"data:image/png;base64,{image_base64}"

def get_responsible_gender_data(responsibles):
    return {
        "Hommes": sum(r.male_count for r in responsibles),
        "Femmes": sum(r.female_count for r in responsibles)
    }



def get_retirement_progression(employee_ages):
    close_to_retirement = sum(1 for emp in employee_ages if emp["age"] is not None and 55 <= emp["age"] < 60)
    return {"Retraite proche": close_to_retirement}

def get_gender_distribution_per_unit():
    results = ResponsableOrganisationUnite.objects.annotate(
        male_count=Count('responsable', filter=Q(responsable__sexe="Masculin")),
        female_count=Count('responsable', filter=Q(responsable__sexe="Féminin"))
    )
    
    # Assurez-vous que les clés sont des chaînes et non des dictionnaires
    return {str(unit.organizational_unit.name): {"Hommes": unit.male_count, "Femmes": unit.female_count} for unit in results}


from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import timedelta
from django.utils.timezone import now
from django.shortcuts import redirect
from datetime import datetime

# Fonction pour obtenir le nom d'utilisateur de la session
def get_username_from_session(request):
    # Implémentation de la récupération du nom d'utilisateur depuis la session
    return request.session.get('username')

def export_to_word(request):
    # Vérifier l'activation
    activation = Activation.objects.first()
    if not activation or not activation.is_valid():
        return redirect("Activation:activation_page")

    username = get_username_from_session(request)
    if not username:
        return redirect('login')

    # Récupérer l'employé correspondant au username de l'utilisateur connecté
    try:
        employee = Employee.objects.get(username=username)
    except Employee.DoesNotExist:
        employee = None  # Si l'employé n'existe pas pour cet utilisateur, gérer le cas ici

    """
    Génère un document Word en fonction des choix de l'utilisateur.
    """
    selected_sections = request.GET.getlist('sections')  # Récupérer les sections choisies
    
    doc = Document()
    section_titles = {
        'employes': 'Liste des Employés',
        'employes_recents': 'Employés récents (moins de 6 mois)',
        'employes_en_conge': 'Employés en congé',
        'employes_revenant_de_conge': 'Employés revenant de congé',
        'employes_depart_retraite': 'Employés en départ pour la retraite (cette année)',
        'employes_retraites': 'Employés déjà retraités',
        'stagiaires': 'Liste des Stagiaires'
    }
    
    # Changer l'orientation du document en paysage
    section = doc.sections[0]
    new_width = section.page_width
    new_height = section.page_height
    section.page_width = new_height
    section.page_height = new_width
    
    doc.add_heading('Rapport de Gestion du Personnel', level=1)

    if 'employes' in selected_sections: 
        doc.add_heading(section_titles['employes'], level=2)

        # Définition des colonnes du tableau
        columns = [
            "Nom Complet", "Matricule", "Sexe", "Date de naissance", "Lieu de naissance", 
            "Prise de service", "Départ à la retraite", "Grade", "Échelon", "Spécialité", 
            "Adresse", "Téléphone", "Email", "Statut"
        ]

        # Création de la table avec en-têtes
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Table Grid'

        # Remplir la première ligne avec les noms des colonnes
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(columns):
            hdr_cells[i].text = col_name

        # Ajouter les données des employés
        for emp in Employee.objects.all():
            row_cells = table.add_row().cells
            row_cells[0].text = f"{emp.last_name} {emp.first_name}"
            row_cells[1].text = emp.matricule if emp.matricule else "N/A"
            row_cells[2].text = emp.sexe if emp.sexe else "N/A"
            row_cells[3].text = emp.date_of_birth.strftime("%d/%m/%Y") if emp.date_of_birth else "N/A"
            row_cells[4].text = emp.place_of_birth if emp.place_of_birth else "N/A"
            row_cells[5].text = emp.start_date.strftime("%d/%m/%Y") if emp.start_date else "N/A"
            row_cells[6].text = emp.retirement_date.strftime("%d/%m/%Y") if emp.retirement_date else "N/A"
            row_cells[7].text = emp.grade if emp.grade else "N/A"
            row_cells[8].text = emp.echelon if emp.echelon else "N/A"
            row_cells[9].text = emp.specialty.designation if emp.specialty else "N/A"
            row_cells[10].text = emp.adresse if emp.adresse else "N/A"
            row_cells[11].text = emp.num_tel if emp.num_tel else "N/A"
            row_cells[12].text = emp.email if emp.email else "N/A"
            row_cells[13].text = emp.status if emp.status else "N/A"

    # Liste des employés récents (moins de 6 mois)
    if 'employes_recents' in selected_sections:
        doc.add_heading(section_titles['employes_recents'], level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nom Complet'
        hdr_cells[1].text = 'Date de début'
        six_mois_avant = now().date() - timedelta(days=180)
        employes_recents = Employee.objects.filter(start_date__gte=six_mois_avant)
        for emp in employes_recents:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{emp.last_name} {emp.first_name}"
            row_cells[1].text = str(emp.start_date)
    
    # Liste des employés en congé
    if 'employes_en_conge' in selected_sections:
        doc.add_heading(section_titles['employes_en_conge'], level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nom Complet'
        hdr_cells[1].text = 'Date de début'
        hdr_cells[2].text = 'Date de fin'
        today = now().date()
        employes_conge = Conge.objects.filter(date_debut__lte=today, date_fin__gte=today)
        for conge in employes_conge:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{conge.employe.first_name} {conge.employe.last_name}"
            row_cells[1].text = str(conge.date_debut)
            row_cells[2].text = str(conge.date_fin)
    
    # Employés revenant de congé
    if 'employes_revenant_de_conge' in selected_sections:
        doc.add_heading(section_titles['employes_revenant_de_conge'], level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nom Complet'
        hdr_cells[1].text = 'Date de début'
        hdr_cells[2].text = 'Date de retour'
        today = now().date()
        employes_revenant = Conge.objects.filter(date_fin__gte=today)
        for conge in employes_revenant:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{conge.employe.last_name} {conge.employe.first_name}"
            row_cells[1].text = str(conge.date_debut)
            row_cells[2].text = str(conge.date_fin)
    
    # Employés en départ pour la retraite cette année
    if 'employes_depart_retraite' in selected_sections:
        doc.add_heading(section_titles['employes_depart_retraite'], level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nom Complet'
        hdr_cells[1].text = 'Date de naissance'
        hdr_cells[2].text = 'Date de départ'
        today = now().date()
        start_of_year = today.replace(month=1, day=1)
        end_of_year = today.replace(month=12, day=31)
        employes_depart_retraite = Employee.objects.filter(retirement_date__gte=start_of_year, retirement_date__lte=end_of_year)
        for emp in employes_depart_retraite:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{emp.first_name} {emp.last_name}"
            row_cells[1].text = str(emp.date_of_birth)
            row_cells[2].text = str(emp.retirement_date)
    
    # Employés déjà retraités
    if 'employes_retraites' in selected_sections:
        doc.add_heading(section_titles['employes_retraites'], level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nom Complet'
        hdr_cells[1].text = 'Date de naissance'
        hdr_cells[2].text = 'Dabte de retraité'
        employes_retraites = Employee.objects.filter(retirement_date__lt=now().date())
        for emp in employes_retraites:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{emp.first_name} {emp.last_name}"
            row_cells[1].text = str(emp.date_of_birth)
            row_cells[2].text = str(emp.retirement_date)
    
    # Liste des stagiaires
    if 'stagiaires' in selected_sections:
        doc.add_heading(section_titles['stagiaires'], level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nom Complet'
        hdr_cells[1].text = 'Date de début'
        stagiaires = Stagiaire.objects.all()
        for stagiaire in stagiaires:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{stagiaire.nom} {stagiaire.prenom}"
            row_cells[1].text = str(stagiaire.date_debut_stage)
    
    # Ajout de "BRAZZAVILLE" + date système en bas à droite du document
    footer_paragraph = doc.add_paragraph()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.add_run("BRAZZAVILLE  " + datetime.now().strftime("%Y-%m-%d")).font.size = Pt(8)

    # Ajouter le nom et prénom de l'employé connecté
    second_footer_paragraph = doc.add_paragraph()
    second_footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if employee:
        second_footer_paragraph.add_run(f"{employee.first_name} {employee.last_name}")
    else:
        second_footer_paragraph.add_run("Nom et prénom de l'employé connecté")

    # Retourner la réponse avec le fichier Word
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=rapport_personnel.docx'
    doc.save(response)
    return response

def export_word_view(request):
    # Vérifier l'activation
    activation = Activation.objects.first()
    if not activation or not activation.is_valid():
        return redirect("Activation:activation_page")
   
    username = get_username_from_session(request)
    if not username:
        return redirect('login')
    """Page d'exportation"""
    return render(request, "export_to_word.html", {'username':username })



